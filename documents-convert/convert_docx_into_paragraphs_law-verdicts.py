from os import listdir, path, makedirs
from docx import Document
from jieba import cut

law_verdicts_train_folder_path = "..\\static\\documents\\docx-law-verdicts-train\\"
law_verdicts_test_folder_path = "..\\static\\documents\\docx-law-verdicts-test\\"
law_verdicts_train_name_list = listdir(law_verdicts_train_folder_path)
law_verdicts_test_name_list = listdir(law_verdicts_test_folder_path)

for file_name in law_verdicts_train_name_list:
    print(file_name)

    document = Document(law_verdicts_train_folder_path + "\\" + file_name)
    paragraphs_text_folder_name = "paragraphs-law-verdicts\\" + file_name[:-5]

    if not path.exists(paragraphs_text_folder_name):
        makedirs(paragraphs_text_folder_name)

    paragraph_index = 1
    for paragraph in document.paragraphs:
        with open(paragraphs_text_folder_name + "\\" + file_name[:-5] + " [Paragraph " + str(paragraph_index) + "].txt", "w", encoding='utf-8') as text_file_object:
            # paragraph_text = " ".join(cut(paragraph.text))
            text_file_object.write(paragraph.text)

        paragraph_index = paragraph_index + 1

for file_name in law_verdicts_test_name_list:
    print(file_name)

    document = Document(law_verdicts_test_folder_path + "\\" + file_name)
    paragraphs_text_folder_name = "paragraphs-law-verdicts\\" + file_name[:-5]

    if not path.exists(paragraphs_text_folder_name):
        makedirs(paragraphs_text_folder_name)

    paragraph_index = 1
    for paragraph in document.paragraphs:
        with open(paragraphs_text_folder_name + "\\" + file_name[:-5] + " [Paragraph " + str(paragraph_index) + "].txt", "w", encoding='utf-8') as text_file_object:
            # paragraph_text = " ".join(cut(paragraph.text))
            text_file_object.write(paragraph.text)

        paragraph_index = paragraph_index + 1